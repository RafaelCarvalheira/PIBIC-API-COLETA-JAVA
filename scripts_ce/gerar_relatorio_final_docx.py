# -*- coding: utf-8 -*-
"""Gera o Relatorio Final de PIBITI a partir do modelo institucional.

Entrada : Modelo_Relatorio_Final (2).docx  (modelo do IME, nao e modificado)
          relatorio_final.xlsx             (resultados computacionais)
          docs/figuras/*.png               (figuras)
Saida   : Relatorio_Final_PIBITI_2025-2026.docx
"""
import os
import re
import statistics

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MODELO = os.path.join(BASE, "Modelo_Relatório_Final (2).docx")
PLANILHA = os.path.join(BASE, "relatorio_final.xlsx")
FIG = os.path.join(BASE, "docs", "figuras")
SAIDA = os.path.join(BASE, "Relatorio_Final_PIBITI_2025-2026.docx")

FONTE = "Times New Roman"
LARG_FIG = Cm(13.5)
GRUPOS = [10, 15, 20, 25, 30]


# ===================================================================== dados
def carregar_dados():
    wb = openpyxl.load_workbook(PLANILHA, data_only=True)

    def linhas(ws, cols):
        out = []
        for r in ws.iter_rows(values_only=True):
            if isinstance(r[0], (int, float)) and r[0]:
                out.append({k: r[i] for k, i in cols.items()})
        return out

    return (
        linhas(wb["CE"], dict(id=0, n=1, og=2, st=3, gap=4, cg=5, oj=9, cj=10)),
        linhas(wb["CEc8"], dict(id=0, n=1, og=2, st=3, gap=4, cg=5, oj=9, cj=10)),
        linhas(wb["CE_A"], dict(id=0, n=1, oa=2, ob=3, og=4, st=5, cg=6, oj=12, cj=13)),
    )


CE, C8, CA = carregar_dados()


def s1(dados):
    return [d for d in dados if d["id"] < 100]


def s2(dados, n=None, faixa=None):
    out = [d for d in dados if d["id"] >= 8001]
    if n is not None:
        out = [d for d in out if d["n"] == n]
    if faixa is not None:
        out = [d for d in out if faixa[0] <= d["id"] <= faixa[1]]
    return out


def dif(d):
    return (d["oj"] - d["og"]) / d["og"] * 100


def med(vals):
    return statistics.mean(vals)


def num(v, casas=2):
    """Formata no padrao brasileiro: 2.294,4 (ponto separa milhar)."""
    if abs(v) < 0.5 * 10 ** (-casas):
        v = 0.0
    return f"{v:,.{casas}f}".replace(",", chr(167)).replace(".", ",").replace(chr(167), ".")


def dif_txt(v):
    if abs(v) < 0.005:
        return "0,00"
    return ("+" if v > 0 else "") + num(v)


# ================================================================ formatacao
def fonte_run(run, tam=12, negrito=None, italico=None):
    run.font.name = FONTE
    run.font.size = Pt(tam)
    if negrito is not None:
        run.bold = negrito
    if italico is not None:
        run.italic = italico


def corpo(p, recuo=True, tam=12, alinh=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.alignment = alinh
    pf.first_line_indent = Cm(1.25) if recuo else Cm(0)
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    for r in p.runs:
        fonte_run(r, tam)
    return p


def celula_limpa(par, tam, negrito=False, alinh=WD_ALIGN_PARAGRAPH.CENTER):
    """Prepara um paragrafo de celula.

    O estilo Normal do modelo tem recuo de primeira linha de 1,25 cm; herdado
    dentro de uma coluna estreita, ele empurra tudo menos o primeiro caractere
    para a linha seguinte. Por isso os recuos sao zerados explicitamente.
    """
    pf = par.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.alignment = alinh
    pf.line_spacing = 1
    pf.space_before = Pt(2 if negrito else 1)
    pf.space_after = Pt(2 if negrito else 1)
    return par


def fixar_larguras(tab, larguras):
    """Impoe larguras de coluna reais: layout fixo, tblGrid e tcW por celula.

    Sem isso o Word redistribui as colunas por autoajuste e quebra os numeros
    no meio (ex.: 279,01 vira "2 / 79,01").
    """
    tbl = tab._tbl
    tab.autofit = False  # grava w:tblLayout type="fixed" na ordem correta
    twips = [int(l.twips) for l in larguras]

    grade = tbl.find(qn("w:tblGrid"))
    if grade is not None:
        colunas = grade.findall(qn("w:gridCol"))
        for col, larg in zip(colunas, twips):
            col.set(qn("w:w"), str(larg))

    for tr in tbl.tr_lst:
        i = 0
        for tc in tr.tc_lst:
            tcPr = tc.get_or_add_tcPr()
            span_el = tcPr.find(qn("w:gridSpan"))
            span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
            largura = sum(twips[i:i + span]) or twips[-1]
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(largura))
            tcW.set(qn("w:type"), "dxa")
            i += span
    return tab


def numerar(par, ilvl, numid=33):
    """Aplica a lista multinivel usada pelo modelo (numId 33) ao paragrafo."""
    pPr = par._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, valor in (("w:ilvl", ilvl), ("w:numId", numid)):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(valor))
        numPr.append(el)
    estilo = pPr.find(qn("w:pStyle"))
    if estilo is not None:
        estilo.addnext(numPr)
    else:
        pPr.insert(0, numPr)
    return par


class Cursor:
    """Insere blocos logo apos um paragrafo-ancora, mantendo a ordem."""

    def __init__(self, doc, ancora):
        self.doc = doc
        self.el = ancora._p

    def _pos(self, obj_el):
        self.el.addnext(obj_el)
        self.el = obj_el

    def p(self, texto="", estilo=None, recuo=True, tam=12,
          alinh=WD_ALIGN_PARAGRAPH.JUSTIFY, negrito=False):
        par = self.doc.add_paragraph(style=estilo)
        if texto:
            escrever_runs(par, texto, tam=tam, negrito=negrito)
        self._pos(par._p)
        if estilo is None:
            corpo(par, recuo=recuo, tam=tam, alinh=alinh)
        return par

    def rico(self, partes, recuo=True, tam=12):
        """partes: lista de (texto, negrito, italico)."""
        par = self.doc.add_paragraph()
        for texto, neg, ital in partes:
            r = par.add_run(texto)
            r.bold, r.italic = neg, ital
        self._pos(par._p)
        return corpo(par, recuo=recuo, tam=tam)

    def titulo(self, texto, nivel=2):
        """Subsecao de terceiro nivel, seguindo a numeracao automatica do modelo."""
        par = self.doc.add_paragraph(style="Heading 1")
        r = par.add_run(texto)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.bold = True
        self._pos(par._p)
        numerar(par, ilvl=2)
        pf = par.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        return par

    def legenda(self, texto, tam=10, antes=0, depois=4):
        par = self.doc.add_paragraph()
        escrever_runs(par, texto, tam=tam)
        self._pos(par._p)
        pf = par.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1
        pf.space_before = Pt(antes)
        pf.space_after = Pt(depois)
        pf.keep_with_next = True
        return par

    def fonte_tabela(self, texto):
        par = self.doc.add_paragraph()
        escrever_runs(par, texto, tam=9)
        self._pos(par._p)
        pf = par.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1
        pf.space_after = Pt(10)
        return par

    def tabela(self, cabecalho, corpo_linhas, larguras=None, tam=9.5,
               alinhamentos=None, mesclar=None, cabecalhos_repetidos=1):
        n_col = len(cabecalho[-1])
        tab = self.doc.add_table(rows=0, cols=n_col)
        tab.style = "Table Grid"
        tab.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._pos(tab._tbl)

        for linha_cab in cabecalho:
            celulas = tab.add_row().cells
            i = 0
            for item in linha_cab:
                if isinstance(item, tuple):
                    texto, span = item
                else:
                    texto, span = item, 1
                alvo = celulas[i] if span == 1 else celulas[i].merge(celulas[i + span - 1])
                par = alvo.paragraphs[0]
                par.text = ""
                escrever_runs(par, texto, tam=tam, negrito=True)
                celula_limpa(par, tam, negrito=True)
                i += span

        for linha in corpo_linhas:
            celulas = tab.add_row().cells
            for j, valor in enumerate(linha):
                par = celulas[j].paragraphs[0]
                par.text = ""
                escrever_runs(par, str(valor), tam=tam)
                al = (alinhamentos[j] if alinhamentos else
                      (WD_ALIGN_PARAGRAPH.LEFT if j == 0 and not str(valor).replace(",", "")
                       .replace(".", "").replace("-", "").isdigit()
                       else WD_ALIGN_PARAGRAPH.CENTER))
                celula_limpa(par, tam, alinh=al)

        if mesclar:
            mesclar_iguais(tab, mesclar)
        if larguras:
            fixar_larguras(tab, larguras)
        manter_junto(tab, cabecalhos_repetidos)
        return tab

    def figura(self, arquivo, largura=LARG_FIG):
        par = self.doc.add_paragraph()
        par.add_run().add_picture(os.path.join(FIG, arquivo), width=largura)
        self._pos(par._p)
        pf = par.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        return par


def def_texto(par, texto, tam=12, negrito=False, alinh=None):
    """Substitui o texto de um paragrafo do modelo preservando o estilo."""
    for r in list(par.runs)[1:]:
        r._element.getparent().remove(r._element)
    if not par.runs:
        par.add_run("")
    r = par.runs[0]
    r.text = texto
    r.bold = negrito
    r.font.size = Pt(tam)
    if alinh is not None:
        par.paragraph_format.alignment = alinh
    return par


def remover(par):
    par._element.getparent().remove(par._element)


# ==================================================================== tabelas
def tab_parametros(cur):
    cur.legenda("Tabela 1 – Parâmetros da meta-heurística por modo de resolução",
                tam=10, antes=8, depois=3)
    # Os parametros que coincidem nos dois modos sao mesclados numa unica
    # celula: o que interessa na tabela e justamente onde os modos divergem.
    cab = [["Parâmetro", "Com simultaneidade", "Sem simultaneidade"]]
    linhas = [
        ["Cliente compartilhado", "um par de tarefas, demandas somadas",
         "um par de tarefas por transportadora"],
        ["Vínculo entre entrega e coleta", "mesmo veículo", "mesma transportadora"],
        ["Alocação do cliente compartilhado", "decidida pela busca (skills)",
         "decidida pela busca (skills)"],
        ["Partidas (multi-start)", "10", "10"],
        ["Iterações por partida", "1.000", "1.000"],
        ["Esforço total (iterações)", "10.000", "10.000"],
        ["Threads", "4", "4"],
        ["Troca de veículo", "habilitada", "habilitada"],
        ["Limiar inicial (θ₀)", "0,03", "0,03"],
        ["Decaimento do limiar (α)", "0,15", "0,15"],
        ["Semente aleatória", "determinística (partida × 1.000 + 42)",
         "determinística (partida × 1.000 + 42)"],
        ["Tamanho da frota", "finita, máx(n, 20) veículos por transportadora",
         "finita, máx(n, 20) veículos por transportadora"],
    ]
    cur.tabela(cab, linhas, larguras=[Cm(4.6), Cm(4.4), Cm(6.0)], tam=9,
               mesclar=(1, 2),
               alinhamentos=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                             WD_ALIGN_PARAGRAPH.CENTER])


def tab_s1(cur):
    cur.legenda("Tabela 2 – Conjunto S1: método exato (Gurobi) e meta-heurística (Jsprit) "
                "no cenário com simultaneidade e colaboração", tam=10, antes=8, depois=3)
    cab = [
        ["", "", ("Gurobi – método exato", 3), ("Jsprit – meta-heurística", 2), ""],
        ["Inst.", "Clientes", "Custo", "Status", "CPU (s)", "Custo", "CPU (s)",
         "Dif. (%)"],
    ]
    linhas = []
    for d in s1(CE):
        linhas.append([
            int(d["id"]), int(d["n"]), num(d["og"]),
            "Ótimo" if d["st"] == "OPTIMAL" else "Lim. tempo",
            num(d["cg"], 1), num(d["oj"]), num(d["cj"], 1),
            dif_txt(dif(d)),
        ])
    linhas.append(["Média", "", "", "", num(med([d["cg"] for d in s1(CE)]), 1), "",
                   num(med([d["cj"] for d in s1(CE)]), 1),
                   dif_txt(med([dif(d) for d in s1(CE)]))])
    cur.tabela(cab, linhas, cabecalhos_repetidos=2, larguras=[Cm(1.4), Cm(1.6), Cm(2.0), Cm(2.2), Cm(2.2),
                                      Cm(2.0), Cm(1.9), Cm(1.7)], tam=9)
    cur.fonte_tabela("Nota: Dif. (%) = (custo Jsprit − custo Gurobi) / custo Gurobi.")


def tab_s2_grupos(cur):
    cur.legenda("Tabela 3 – Conjunto S2: resultados médios por número de clientes "
                "no cenário com simultaneidade e colaboração", tam=10, antes=8, depois=3)
    cab = [
        ["", ("Gurobi – método exato", 3), ("Jsprit – meta-heurística", 2),
         ("Comparação", 2)],
        ["Clientes", "Custo", "CPU (s)", "Ótimas", "Custo", "CPU (s)", "Dif. (%)",
         "Acel."],
    ]
    linhas = []
    for n in GRUPOS:
        sub = s2(CE, n)
        cg, cj = med([d["cg"] for d in sub]), med([d["cj"] for d in sub])
        linhas.append([
            n, num(med([d["og"] for d in sub])), num(cg, 1),
            f"{sum(1 for d in sub if d['st'] == 'OPTIMAL')}/20",
            num(med([d["oj"] for d in sub])), num(cj, 1),
            dif_txt(med([dif(d) for d in sub])),
            num(cg / cj, 1) + "×",
        ])
    tot = s2(CE)
    cgt, cjt = med([d["cg"] for d in tot]), med([d["cj"] for d in tot])
    linhas.append(["Total", "—", num(cgt, 1),
                   f"{sum(1 for d in tot if d['st'] == 'OPTIMAL')}/100", "—", num(cjt, 1),
                   dif_txt(med([dif(d) for d in tot])), num(cgt / cjt, 1) + "×"])
    cur.tabela(cab, linhas, cabecalhos_repetidos=2, larguras=[Cm(1.9), Cm(2.1), Cm(2.1), Cm(1.7), Cm(2.1),
                                      Cm(2.1), Cm(1.6), Cm(1.4)], tam=9)
    cur.fonte_tabela("Nota: a instância 8090 foi executada sem limite de tempo "
                     "(39.532,9 s); desconsiderando-a, o tempo médio do grupo de 20 clientes "
                     "cai para 1.097,7 s, ainda cerca de 34 vezes o tempo do Jsprit.")


def tab_random_clustered(cur):
    cur.legenda("Tabela 4 – Conjunto S2: desempenho nos subconjuntos aleatório (S2R) "
                "e agrupado (S2C)", tam=10, antes=8, depois=3)
    cab = [
        ["", ("S2R – aleatórias (50)", 3), ("S2C – agrupadas (50)", 3)],
        ["Cenário", "Dif. (%)", "Iguais", "Acel.", "Dif. (%)", "Iguais", "Acel."],
    ]
    cenarios = [("Com simultaneidade", CE), ("Sem simultaneidade", C8),
                ("Sem compartilhamento", CA)]
    linhas = []
    for rotulo, dados in cenarios:
        linha = [rotulo]
        for faixa in [(8001, 8050), (8051, 8100)]:
            sub = s2(dados, faixa=faixa)
            difs = [dif(d) for d in sub]
            acel = med([d["cg"] for d in sub]) / med([d["cj"] for d in sub])
            linha += [dif_txt(med(difs)),
                      str(sum(1 for x in difs if abs(x) < 1e-6)),
                      num(acel, 1) + "×"]
        linhas.append(linha)
    cur.tabela(cab, linhas, cabecalhos_repetidos=2, larguras=[Cm(4.2), Cm(1.8), Cm(1.6), Cm(1.6),
                                      Cm(1.8), Cm(1.6), Cm(1.6)], tam=9,
               alinhamentos=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6)
    cur.fonte_tabela("Nota: “Iguais” indica o número de instâncias em que o Jsprit "
                     "reproduziu exatamente o custo do método exato.")


def tab_cenarios(cur):
    cur.legenda("Tabela 5 – Diferença percentual média por cenário e por número de clientes "
                "(conjunto S2)", tam=10, antes=8, depois=3)
    cab = [["Cenário"] + [f"{n} cli." for n in GRUPOS] + ["Média", "Iguais"]]
    linhas = []
    for rotulo, dados in [("Com simultaneidade (colaborativo)", CE),
                          ("Sem simultaneidade (colaborativo)", C8),
                          ("Sem compartilhamento", CA)]:
        linha = [rotulo]
        for n in GRUPOS:
            v = med([dif(d) for d in s2(dados, n)])
            linha.append(dif_txt(v))
        tot = [dif(d) for d in s2(dados)]
        linha.append(dif_txt(med(tot)))
        linha.append(f"{sum(1 for x in tot if abs(x) < 1e-6)}/100")
        linhas.append(linha)
    cur.tabela(cab, linhas, larguras=[Cm(5.0)] + [Cm(1.35)] * 5 + [Cm(1.5), Cm(1.6)], tam=9,
               alinhamentos=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 7)


def tab_colaboracao(cur):
    cur.legenda("Tabela 6 – Economia proporcionada pela colaboração horizontal: "
                "método exato e meta-heurística", tam=10, antes=8, depois=3)
    mapa = {d["id"]: d for d in CA}

    def economias(sel):
        sub = [d for d in CE if sel(d["id"]) and d["id"] in mapa]
        eg = med([(mapa[d["id"]]["og"] - d["og"]) / mapa[d["id"]]["og"] * 100 for d in sub])
        ej = med([(mapa[d["id"]]["oj"] - d["oj"]) / mapa[d["id"]]["oj"] * 100 for d in sub])
        return eg, ej, len(sub)

    cab = [
        ["", "", ("Economia da colaboração (%)", 2), ""],
        ["Conjunto", "Instâncias", "Gurobi", "Jsprit", "Dif. (p.p.)"],
    ]
    linhas = []
    grupos = [("S1 (12 instâncias)", lambda i: i < 100)]
    for n in GRUPOS:
        ids = {d["id"] for d in s2(CE, n)}
        grupos.append((f"S2 – {n} clientes", lambda i, ids=ids: i in ids))
    grupos += [("S2R – aleatórias", lambda i: 8001 <= i <= 8050),
               ("S2C – agrupadas", lambda i: 8051 <= i <= 8100),
               ("S2 – total", lambda i: i >= 8001)]
    for rotulo, sel in grupos:
        eg, ej, k = economias(sel)
        linhas.append([rotulo, k, num(eg), num(ej), num(abs(eg - ej))])
    cur.tabela(cab, linhas, cabecalhos_repetidos=2, larguras=[Cm(4.5), Cm(2.2), Cm(2.6), Cm(2.6), Cm(2.4)], tam=9,
               alinhamentos=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4)
    cur.fonte_tabela("Nota: economia = (custo sem compartilhamento − custo colaborativo) / "
                     "custo sem compartilhamento, calculada instância a instância e depois "
                     "promediada.")


# ===================================================================== texto
# Termos estrangeiros grafados em italico. Nomes proprios de produtos e
# linguagens (Jsprit, Gurobi, Java, Spring Boot, Julia, JuMP, Python) e siglas
# (API, REST, JSON, CSV) ficam em redondo, conforme o uso corrente.
TERMOS_ESTRANGEIROS = [
    "Shared Customer Collaboration Vehicle Routing Problem with Simultaneous "
    "Pickup and Delivery",
    "European Journal of Operational Research",
    "Vehicle Routing Problem",
    "Ruin-and-Recreate",
    "Threshold Accepting",
    "SameVehicleConstraint",
    "VrpService.java",
    "multi-start",
    "framework",
    "endpoints",
    "endpoint",
    "software",
    "skills",
    "threads",
    "gap",
]

_PADRAO_ITALICO = re.compile(
    "(" + "|".join(re.escape(t) for t in
                   sorted(TERMOS_ESTRANGEIROS, key=len, reverse=True)) + ")"
)
_CONJUNTO_ITALICO = set(TERMOS_ESTRANGEIROS)


def escrever_runs(par, texto, tam=12, negrito=False):
    """Escreve o texto quebrando-o em runs, com italico nos termos estrangeiros."""
    for pedaco in _PADRAO_ITALICO.split(texto):
        if not pedaco:
            continue
        run = par.add_run(pedaco)
        run.bold = negrito
        run.italic = pedaco in _CONJUNTO_ITALICO
        fonte_run(run, tam)
    if not par.runs:
        fonte_run(par.add_run(""), tam)
    return par


def mesclar_iguais(tab, colunas):
    """Funde celulas vizinhas de mesmo conteudo nas colunas indicadas.

    Evita repetir o mesmo valor lado a lado, que era o que deixava a tabela de
    parametros visualmente poluida.
    """
    ini, fim = colunas
    for linha in tab.rows[1:]:
        celulas = linha.cells
        if len(celulas) <= fim:
            continue
        textos = [celulas[j].text.strip() for j in range(ini, fim + 1)]
        if len(set(textos)) == 1 and textos[0]:
            for j in range(ini + 1, fim + 1):
                celulas[j].text = ""
            alvo = celulas[ini].merge(celulas[fim])
            # a fusao conserva os paragrafos vazios das celulas absorvidas, o
            # que inchava a altura da linha; sobra apenas o primeiro
            for extra in alvo.paragraphs[1:]:
                extra._element.getparent().remove(extra._element)
    return tab


def manter_junto(tab, repetir_cabecalho=1):
    """Impede que a tabela se parta entre paginas e repete o cabecalho."""
    tbl = tab._tbl
    linhas = tbl.tr_lst
    for i, tr in enumerate(linhas):
        trPr = tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        if i < repetir_cabecalho:
            trPr.append(OxmlElement("w:tblHeader"))
    for tr in linhas[:-1]:
        for par in tr.iter(qn("w:p")):
            pPr = par.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                par.insert(0, pPr)
            pPr.append(OxmlElement("w:keepNext"))
    return tab


def cor_vermelha(valor):
    if not valor or len(valor) != 6:
        return False
    try:
        r, g, b = (int(valor[i:i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return False
    return r > 120 and r > g + 50 and r > b + 50


def limpar_runs_vermelhos(elemento):
    """Remove, no nivel XML, todo run cuja cor explicita seja avermelhada."""
    alvos = []
    for run in elemento.iter(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            continue
        cor = rPr.find(qn("w:color"))
        if cor is not None and cor_vermelha((cor.get(qn("w:val")) or "").upper()):
            alvos.append(run)
    for run in alvos:
        run.getparent().remove(run)
    return len(alvos)


def remover_orientacoes(doc):
    """Apaga o texto em vermelho do modelo (instrucoes de preenchimento).

    Percorre o corpo e TODAS as partes de cabecalho/rodape do pacote. A API de
    secoes do python-docx nao alcanca todos os footerN.xml quando o documento
    tem secoes com rodapes distintos — era por onde escapava o rodape
    "(O RELATORIO DEVE CONTER NO MAXIMO 25 PAGINAS)". O campo de numero de
    pagina nao tem cor explicita e por isso e preservado.
    """
    removidos = limpar_runs_vermelhos(doc.element.body)
    for rel in doc.part.rels.values():
        if "header" in rel.reltype or "footer" in rel.reltype:
            removidos += limpar_runs_vermelhos(rel.target_part.element)
    return removidos


def preencher_capa(doc):
    p = doc.paragraphs
    def_texto(p[9], "Programa Institucional de Bolsas de Iniciação em Desenvolvimento "
                    "Tecnológico e Inovação (PIBITI)")
    def_texto(p[10], "Edital 2025/2026")
    def_texto(p[13], "Adaptação e comparação de algoritmos para o problema de roteamento de veículos utilizando programação em Julia")
    def_texto(p[15], "Projeto: Pesquisa Operacional: Aplicação na Logística Militar")
    def_texto(p[23], "Julho de 2026")
    def_texto(p[27], "Autor: Rafael Vargas Carvalheira")
    def_texto(p[28], "E-mail: rafaelvargascar20@gmail.com")
    def_texto(p[29], "Telefone: (12) 98141-8480")
    def_texto(p[31], "Orientador: Prof. Orivalde Soares da Silva Júnior – D.Sc.")
    def_texto(p[32], "E-mail: orivalde@ime.eb.br")
    def_texto(p[33], "Telefone: (21) 99939-8145")
    for i in (37, 36, 35):  # co-orientador nao se aplica
        remover(p[i])


def escrever_resumo(cur):
    cur.p("A colaboração horizontal entre transportadoras que atendem clientes em comum "
          "permite reduzir de forma expressiva o custo agregado de transporte. Este "
          "trabalho trata do Shared Customer Collaboration Vehicle Routing Problem with "
          "Simultaneous Pickup and Delivery (SCCVRPSPD), extensão do problema de "
          "roteamento colaborativo com clientes compartilhados na qual cada cliente "
          "possui, simultaneamente, demanda de entrega e de coleta. O modelo exato, "
          "formulado em programação linear inteira mista e resolvido com o otimizador Gurobi, "
          "não converge dentro de limites de tempo praticáveis à medida que o número de "
          "clientes cresce, o que inviabiliza seu uso operacional. Desenvolveu-se, "
          "portanto, um serviço web em Java (Spring Boot) que expõe uma meta-heurística "
          "baseada no esquema Ruin-and-Recreate, implementada com a biblioteca Jsprit, "
          "capaz de resolver quatro cenários do problema: com e sem a restrição de "
          "simultaneidade, cada qual com e sem colaboração entre as transportadoras. "
          "A validação empregou 112 instâncias adaptadas da literatura, previamente "
          "resolvidas pelo método exato. No cenário completo, a meta-heurística apresentou "
          "diferença média de custo de 0,22% no conjunto de 12 instâncias e de 0,36% no "
          "conjunto de 100 instâncias, reproduzindo exatamente a solução ótima em 60 dos "
          "100 casos e superando o incumbente do otimizador em instâncias não resolvidas até "
          "a otimalidade. Nesse cenário o tempo de processamento manteve-se abaixo de 82 "
          "segundos em todas as execuções, contra médias superiores a 6.500 segundos do "
          "método exato nas maiores instâncias. Verificou-se ainda que a meta-heurística estima a "
          "economia proporcionada pela colaboração, de 9,50% em média, com desvio "
          "inferior a 0,2 ponto percentual em relação ao método exato, o que a habilita "
          "como ferramenta prática de apoio à decisão.", recuo=False)
    cur.rico([("Palavras-chave: ", True, False),
              ("roteamento de veículos; colaboração horizontal; coleta e entrega "
               "simultâneas; meta-heurística; Ruin-and-Recreate.", False, False)],
             recuo=False)


def escrever_introducao(cur):
    cur.p("O transporte rodoviário de cargas fracionadas em áreas urbanas é realizado, "
          "em geral, por diversas empresas que operam simultaneamente sobre a mesma "
          "malha viária, partindo de centros de distribuição distintos e atendendo, com "
          "frequência, clientes localizados nas mesmas regiões. Nesse contexto, o "
          "Problema de Roteamento de Veículos (VRP, do inglês Vehicle Routing Problem) "
          "ocupa posição central na literatura de otimização combinatória, tanto por sua "
          "relevância econômica quanto por sua dificuldade computacional, uma vez que "
          "pertence à classe dos problemas NP-difíceis.")
    cur.p("A colaboração horizontal entre transportadoras surge como estratégia "
          "promissora de redução de custos. Quando duas empresas atendem clientes "
          "coincidentes ou muito próximos, transferir o atendimento de um cliente para a "
          "rota mais favorável da outra companhia encurta distâncias, melhora a "
          "utilização da capacidade da frota e reduz o número de veículos necessários "
          "(PAN et al., 2019; GANSTERER; HARTL, 2018). Essa vertente foi formalizada por "
          "Fernández, Roca-Riu e Speranza (2018) sob a denominação Shared Customer "
          "Collaboration Vehicle Routing Problem (SCCVRP), no qual clientes ditos "
          "compartilhados possuem demandas junto a mais de uma transportadora e podem ser "
          "atendidos por qualquer uma delas. Extensões posteriores incorporaram janelas "
          "de tempo e limites de custo por transportadora (HIMSTEDT; MEISEL, 2021), "
          "objetivos múltiplos (ZHANG et al., 2020) e abordagens meta-heurísticas "
          "híbridas (TORRES-RAMOS et al., 2019).")
    cur.p("Uma segunda variante relevante do VRP é a de coleta e entrega, na qual os "
          "clientes tanto recebem quanto devolvem mercadorias (PLOSKAS et al., 2015). "
          "Sua importância cresce com a consolidação da logística reversa e com a "
          "necessidade de aproveitar o trajeto de retorno dos veículos, evitando "
          "quilometragem ociosa (RAN; LI; ZHAO, 2022). Quando a coleta e a entrega devem "
          "ocorrer em uma única visita ao cliente, tem-se a variante simultânea "
          "(CACERES-CRUZ et al., 2014; PADMANABHAN et al., 2022), que impõe ao algoritmo "
          "o controle da carga flutuante do veículo ao longo de toda a rota.")
    cur.p("A união dessas duas vertentes dá origem ao problema tratado neste trabalho, "
          "denominado SCCVRPSPD, no qual transportadoras distintas podem atender "
          "conjuntamente clientes compartilhados que possuem, ao mesmo tempo, demandas de "
          "entrega e de coleta, com a exigência de que cada cliente seja visitado uma "
          "única vez. O problema é diretamente aplicável à logística militar, domínio em "
          "que diferentes unidades operam depósitos próprios e atendem organizações "
          "militares comuns, e no qual a coordenação do suprimento pode gerar economia "
          "significativa de recursos.")
    cur.p("Do ponto de vista computacional, a formulação exata do problema, baseada em "
          "carga e resolvida por programação linear inteira mista, garante otimalidade "
          "apenas para instâncias de pequeno porte. Para instâncias maiores, o esforço "
          "computacional cresce de forma acentuada e o otimizador frequentemente encerra a "
          "execução por limite de tempo, sem comprovação de otimalidade. Abordagens "
          "meta-heurísticas constituem, nesse cenário, a alternativa tratável. Entre "
          "elas, o esquema Ruin-and-Recreate (SCHRIMPF et al., 2000) destaca-se pela "
          "simplicidade conceitual e pela qualidade das soluções obtidas, estando "
          "disponível em implementação madura na biblioteca Jsprit.")
    cur.p("O presente relatório descreve o desenvolvimento de um serviço web que "
          "disponibiliza algoritmos de roteirização meta-heurísticos para o SCCVRPSPD e "
          "a validação sistemática de suas soluções contra os resultados do método "
          "exato. O texto está organizado como segue: a Justificativa expõe o problema "
          "que motiva a pesquisa; os Objetivos delimitam o escopo; a Metodologia "
          "apresenta o modelo de referência, a arquitetura do serviço, a meta-heurística "
          "e o protocolo experimental; a seção de Resultados e Análise reporta os "
          "experimentos conduzidos sobre 112 instâncias; e a Conclusão sintetiza as "
          "contribuições e aponta desdobramentos.")


def escrever_justificativa(cur):
    cur.p("O emprego de métodos exatos para o SCCVRPSPD encontra um limite prático bem "
          "definido. Nos experimentos conduzidos neste trabalho, o tempo médio de "
          "processamento do otimizador Gurobi passou de 2,7 segundos, para instâncias com 10 "
          "clientes, a 6.542,1 segundos para instâncias com 30 clientes, um crescimento de "
          "mais de três ordens de grandeza para uma triplicação do porte do problema. "
          "Ainda assim, apenas 2 das 20 instâncias com 30 clientes foram resolvidas até a "
          "otimalidade comprovada dentro do limite de 7.200 segundos, e 31 das 100 "
          "instâncias do conjunto ampliado foram encerradas com gap em aberto.")
    cur.p("Esse comportamento inviabiliza a aplicação direta do método exato em sistemas "
          "de apoio à decisão. Um planejador logístico que precise avaliar alternativas "
          "de roteirização, simular cenários de colaboração ou reagir a alterações de "
          "demanda não dispõe de duas horas de processamento por consulta. A demanda "
          "operacional é por respostas em segundos, com qualidade de solução "
          "suficientemente próxima do ótimo para que a decisão tomada não seja "
          "materialmente diferente daquela que o método exato indicaria.")
    cur.p("Há, portanto, um problema de pesquisa bem delimitado: verificar se uma "
          "meta-heurística consegue, para este problema específico, entregar soluções "
          "cuja distância ao ótimo seja pequena e mensurável, em tempo compatível com uso "
          "interativo. A verificação exige uma referência confiável, isto é, soluções "
          "ótimas comprovadas para um conjunto de instâncias suficientemente amplo, e "
          "um protocolo de comparação uniforme, aplicado às mesmas instâncias e sob os "
          "mesmos parâmetros de modelagem.")
    cur.p("Justifica-se ainda a construção de um serviço web, e não de um programa "
          "isolado. A disponibilização dos algoritmos por meio de uma interface de "
          "programação de aplicações desacopla o motor de resolução da linguagem e do ambiente do "
          "cliente, permite que o mesmo motor atenda diferentes sistemas e viabiliza a "
          "condução automatizada de baterias experimentais, recurso decisivo para a "
          "execução das 336 resoluções que compõem os experimentos aqui reportados.")


def escrever_objetivos(cur):
    cur.titulo("Objetivo geral", nivel=2)
    cur.p("Desenvolver e validar uma infraestrutura de software, na forma de serviço web, "
          "que resolva o problema de roteamento de veículos colaborativo com clientes "
          "compartilhados e coletas e entregas simultâneas por meio de algoritmos "
          "meta-heurísticos, tomando como referência de qualidade as soluções ótimas "
          "obtidas por método exato.")
    cur.titulo("Objetivos específicos", nivel=2)
    itens = [
        "Projetar e implementar uma interface de programação de aplicações (API REST) "
        "capaz de receber instâncias de roteamento, processá-las e retornar as rotas "
        "otimizadas em formato estruturado;",
        "Modelar computacionalmente, no framework Jsprit, os quatro cenários do "
        "problema: com e sem a restrição de simultaneidade, cada um deles com e sem "
        "colaboração horizontal entre as transportadoras;",
        "Conduzir a validação experimental das implementações sobre os conjuntos de "
        "instâncias S1 (12 instâncias) e S2 (100 instâncias), comparando custo da função "
        "objetivo e tempo de processamento com os resultados do otimizador Gurobi;",
        "Quantificar, por meio da meta-heurística, a economia proporcionada pela "
        "colaboração horizontal e avaliar a fidelidade dessa estimativa frente à obtida "
        "pelo método exato.",
    ]
    for i, texto in enumerate(itens, 1):
        p = cur.p(f"{i}. {texto}", recuo=False)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)


def escrever_metodologia(cur):
    cur.p("A metodologia adotada compreende cinco etapas encadeadas: a especificação do "
          "modelo de referência, a construção do serviço web, a configuração da "
          "meta-heurística, a tradução das restrições do modelo matemático para o "
          "paradigma do framework Jsprit e, por fim, a definição do protocolo "
          "experimental de comparação.")

    cur.titulo("O problema SCCVRPSPD e o modelo de referência", nivel=2)
    cur.p("Considera-se um conjunto de transportadoras, cada uma operando um único "
          "depósito com frota de capacidade homogênea, e um conjunto de clientes "
          "distribuídos em uma região. Cada cliente possui, para cada transportadora à "
          "qual encomendou serviço, uma quantidade de mercadorias a ser entregue e uma "
          "quantidade a ser coletada. Clientes que possuem demanda junto a mais de uma "
          "transportadora são denominados compartilhados e podem ser atendidos por "
          "qualquer uma delas. O objetivo é minimizar o custo total das rotas de todas as "
          "transportadoras.")
    cur.p("O modelo de referência é uma formulação de programação linear inteira mista "
          "baseada em carga, derivada de Fernández, Roca-Riu e Speranza (2018) e "
          "estendida com as demandas de coleta. Variáveis binárias de roteamento indicam "
          "o percurso de cada arco por cada transportadora; variáveis binárias de "
          "alocação indicam qual transportadora atende a demanda de cada cliente; e "
          "variáveis contínuas de fluxo descrevem as cargas de entrega e de coleta "
          "transportadas em cada arco, o que assegura simultaneamente o atendimento "
          "correto das demandas, o respeito à capacidade dos veículos e a eliminação de "
          "sub-rotas.")
    cur.p("Duas restrições concentram o interesse deste trabalho. A primeira estabelece "
          "que cada cliente seja visitado exatamente uma vez pelo conjunto das "
          "transportadoras. É ela que impõe a simultaneidade da coleta e da entrega, "
          "obrigando ambas as operações a ocorrerem em uma única parada. A segunda "
          "governa a colaboração: quando desativada, cada transportadora atende apenas os "
          "clientes de sua própria carteira, e o cenário degenera em dois subproblemas "
          "independentes cujos custos são somados. A combinação dessas duas dimensões "
          "define os quatro cenários implementados no serviço.")
    cur.p("O modelo exato foi implementado em linguagem Julia, com o pacote JuMP e o "
          "otimizador Gurobi, em trabalho anterior do grupo de pesquisa. Seus resultados "
          "constituem a referência de qualidade contra a qual a meta-heurística é "
          "avaliada, e não são contribuição deste subprojeto.")

    cur.titulo("Arquitetura do serviço web", nivel=2)
    cur.p("O serviço foi implementado em Java 17 sobre o framework Spring Boot 3.2, "
          "seguindo a organização em camadas Controller–Service–Solver. A camada de "
          "controle expõe quatro endpoints REST, um para cada cenário do problema, que "
          "recebem a instância em formato JSON (parâmetros globais, frotas e depósitos, "
          "clientes com suas demandas de entrega e coleta por transportadora, e a matriz "
          "de custos) e devolvem a solução com o custo total, a sequência de atividades "
          "de cada rota e o custo por rota. A camada de serviço constrói o objeto de "
          "problema do Jsprit, configura a matriz de custos, executa o algoritmo e mapeia "
          "a solução para o formato de resposta.")
    cur.p("Em torno do serviço foi construído um pipeline experimental em Python, "
          "responsável por converter as instâncias no formato original para JSON, "
          "submeter automaticamente todas as instâncias aos quatro endpoints, coletar as "
          "respostas e consolidar os resultados em planilhas comparativas. A Figura 1 "
          "resume a arquitetura e o fluxo de dados, incluindo o ramo do método exato que "
          "fornece os valores de referência.")
    cur.figura("fig1_arquitetura.png")
    cur.legenda("Figura 1 – Arquitetura do serviço e fluxo experimental", tam=10, depois=10)

    cur.titulo("A meta-heurística Ruin-and-Recreate", nivel=2)
    cur.p("O motor de resolução é a biblioteca Jsprit 1.9.0-beta.3, que implementa o "
          "esquema Ruin-and-Recreate proposto por Schrimpf et al. (2000). O algoritmo "
          "parte de uma solução inicial construída por inserção com arrependimento e "
          "alterna iterativamente duas fases. Na fase de destruição, um subconjunto dos "
          "clientes é removido das rotas, seja pela estratégia radial, que remove "
          "clientes próximos a um ponto de referência sorteado, seja pela estratégia "
          "aleatória. Na fase de reconstrução, os clientes removidos são reinseridos pela "
          "heurística de melhor inserção, que testa todas as posições factíveis e escolhe "
          "a de menor acréscimo de custo, ou pela inserção com arrependimento, que "
          "prioriza o cliente cuja segunda melhor posição é sensivelmente pior que a "
          "melhor, reduzindo o risco de inserções míopes.")
    cur.p("A solução resultante de cada iteração é submetida a um critério de aceitação "
          "do tipo Threshold Accepting, variante determinística do recozimento simulado, "
          "que admite pioras limitadas por um limiar decrescente ao longo da execução. "
          "Para mitigar a convergência para mínimos locais, adotou-se uma estratégia de "
          "múltiplos pontos de partida com sementes determinísticas: cada partida "
          "reinstancia o algoritmo e executa mil iterações internas, retendo-se ao final "
          "a melhor solução encontrada. A Figura 2 esquematiza o procedimento e a "
          "Tabela 1 consolida os parâmetros efetivamente empregados.")
    cur.figura("fig2_ruin_recreate.png")
    cur.legenda("Figura 2 – Esquema do algoritmo Ruin-and-Recreate com múltiplas partidas",
                tam=10, depois=10)
    tab_parametros(cur)

    cur.titulo("Tradução das restrições para o Jsprit", nivel=2)
    cur.p("A principal dificuldade metodológica do trabalho consistiu em traduzir "
          "restrições formuladas de modo declarativo, por meio de inequações, para o "
          "paradigma orientado a objetos do Jsprit, no qual elas devem ser expressas por "
          "classes de estado e funções de atualização executadas a cada modificação de "
          "rota.")
    cur.p("As demandas de coleta e entrega de cada cliente foram modeladas como duas "
          "tarefas distintas: uma de entrega, que reduz a carga do veículo, e uma de "
          "coleta, que a aumenta. Essa separação é necessária para que a verificação de "
          "capacidade seja feita corretamente em cada ponto da rota, respeitando a carga "
          "acumulada ao longo do trajeto. Como o Jsprit não oferece nativamente vínculo "
          "entre tarefas independentes, o par foi amarrado por uma restrição customizada "
          "de rota rígida, apoiada em um atualizador de estado que rastreia, a cada "
          "modificação, qual veículo atende cada tarefa.")
    cur.p("A alocação dos clientes compartilhados não é fixada de antemão em nenhum dos "
          "cenários: ela é deixada a cargo da própria busca, exatamente como a variável "
          "binária de alocação do modelo matemático. Isso é obtido pelo mecanismo nativo "
          "de habilidades do Jsprit. Cada veículo recebe a habilidade correspondente à "
          "sua transportadora e as tarefas de um cliente exclusivo exigem a habilidade "
          "do seu atendente, ao passo que as tarefas de um cliente compartilhado não "
          "recebem exigência alguma, de modo que o algoritmo decide, ao inseri-las, qual "
          "transportadora as atende.")
    cur.p("A distinção entre os dois cenários está na representação do cliente "
          "compartilhado e na natureza do vínculo. No cenário com simultaneidade, as "
          "demandas das duas transportadoras são somadas em um único par de tarefas, e a "
          "restrição exige que entrega e coleta sejam atendidas pelo mesmo veículo, o "
          "que reproduz a visita única imposta pelo modelo. No cenário sem "
          "simultaneidade, o cliente recebe um par de tarefas por transportadora, e a "
          "restrição exige apenas que entrega e coleta de uma mesma demanda fiquem com a "
          "mesma transportadora, sem obrigá-las ao mesmo veículo. Essa formulação "
          "reproduz o modelo de referência, no qual as variáveis de roteamento são "
          "indexadas por transportadora, sem índice de veículo, e uma única variável de "
          "alocação governa entrega e coleta. Ela admite, em consequência, tanto que as "
          "duas transportadoras atendam separadamente o mesmo cliente quanto que uma "
          "delas o visite em duas passagens distintas.")

    cur.titulo("Instâncias e protocolo experimental", nivel=2)
    cur.p("Os experimentos empregaram os dois conjuntos de instâncias utilizados por "
          "Fernández, Roca-Riu e Speranza (2018), adaptados com a inclusão de demandas de "
          "coleta. O conjunto S1 reúne 12 instâncias com 18 a 30 clientes. O conjunto S2 "
          "reúne 100 instâncias, sendo 50 com clientes distribuídos aleatoriamente em um "
          "quadrado de 100 unidades de lado (subconjunto S2R) e 50 com clientes agrupados "
          "em três a cinco aglomerados (subconjunto S2C); em cada subconjunto há 10 "
          "instâncias para cada porte de 10, 15, 20, 25 e 30 clientes. As demandas de "
          "coleta foram geradas de modo que sua soma, por transportadora, igualasse a "
          "soma das entregas, preservando a coerência do problema original e evitando "
          "alterações na capacidade dos veículos.")
    cur.p("A avaliação abrange os três cenários para os quais o método exato dispõe de "
          "rotina de resolução equivalente, o que permite comparar cada solução com uma "
          "referência: os cenários com e sem simultaneidade em modo colaborativo e o "
          "cenário sem compartilhamento. As execuções ocorreram em processador Intel Core i7 de "
          "3,20 GHz com 16 GB de memória RAM, mesma configuração empregada nas execuções "
          "do método exato, o que torna diretamente comparáveis os tempos reportados. Ao "
          "otimizador Gurobi foi imposto limite de 7.200 segundos por execução; à "
          "meta-heurística não se impôs limite de tempo, sendo o critério de parada "
          "exclusivamente o número de iterações.")
    cur.p("A métrica de qualidade adotada é a diferença percentual entre o custo obtido "
          "pela meta-heurística e o custo obtido pelo método exato, calculada instância a "
          "instância. Valores positivos indicam solução mais cara que a de referência; "
          "valores negativos indicam que a meta-heurística encontrou solução melhor que o "
          "incumbente do otimizador, situação possível quando este encerra a execução por "
          "limite de tempo sem comprovar otimalidade. A métrica de eficiência é o tempo "
          "de processamento, reportado em segundos e agregado pela média dos grupos.")


def escrever_resultados(cur):
    cur.p("Esta seção apresenta os resultados das 336 execuções comparadas com a "
          "referência exata, correspondentes às 112 instâncias resolvidas em cada um dos "
          "três cenários avaliados. A "
          "apresentação segue a ordem: validação no conjunto reduzido, escalabilidade no "
          "conjunto ampliado, efeito da distribuição espacial dos clientes, comportamento "
          "nos demais cenários, economia proporcionada pela colaboração e, por fim, "
          "discussão das limitações identificadas.")

    cur.titulo("Validação no conjunto S1", nivel=2)
    cur.p("A Tabela 2 apresenta, para cada uma das 12 instâncias do conjunto S1, os "
          "resultados do método exato e da meta-heurística no cenário completo do "
          "problema, com simultaneidade e com colaboração.")
    tab_s1(cur)
    cur.p("A diferença percentual média foi de +0,22%. Em seis instâncias a "
          "meta-heurística reproduziu exatamente o custo ótimo e, na instância 10, a "
          "diferença limitou-se a uma centésima de unidade de custo, em favor do "
          "Jsprit. O resultado mais expressivo ocorreu na instância 8, cujo "
          "processamento exato foi interrompido por limite de tempo com 1,79% de gap "
          "em aberto: a meta-heurística obteve custo de 243,60 contra 248,10 do "
          "incumbente, uma redução de 1,81%. O pior desempenho relativo ocorreu na "
          "instância 2, com desvio de +2,69%, também ela encerrada por limite de "
          "tempo pelo método exato.")
    cur.p("Quanto ao tempo, o contraste é expressivo. O método exato consumiu, em média, "
          "2.294,4 segundos por instância, com três execuções encerradas no limite de "
          "7.200 segundos, enquanto a meta-heurística concluiu todas as execuções em "
          "média de 41,1 segundos e máximo de 76,9 segundos, o que representa aceleração média de "
          "55,8 vezes, sem limite de tempo imposto.")

    cur.titulo("Escalabilidade no conjunto S2", nivel=2)
    cur.p("O conjunto S2, com 100 instâncias, permite observar como a diferença de "
          "qualidade e o custo computacional evoluem com o porte do problema. A Tabela 3 "
          "consolida os resultados médios por número de clientes.")
    tab_s2_grupos(cur)
    cur.p("Três padrões emergem dos dados. O primeiro é a degradação controlada da "
          "qualidade: a diferença média parte de 0,00% no grupo de 10 clientes, mantém-se "
          "ligeiramente negativa no grupo de 15 clientes, isto é, a meta-heurística foi, "
          "em média, marginalmente melhor que o otimizador, e cresce de forma monotônica até "
          "+1,17% no grupo de 30 clientes, permanecendo em +0,36% na média geral. O "
          "segundo é a fidelidade das soluções: em 60 das 100 instâncias o custo obtido "
          "coincide exatamente com o do método exato, e em 4 instâncias o supera. O "
          "terceiro é a inversão do regime de tempo: para 10 clientes a meta-heurística é "
          "mais lenta que o otimizador, situação que se inverte a partir de 15 clientes e se "
          "acentua rapidamente, atingindo aceleração de 96,9 vezes no grupo de 30 "
          "clientes.")
    cur.p("Cabe observar que parte da diferença medida nos grupos maiores não decorre de "
          "deficiência da meta-heurística, mas da própria imprecisão da referência: no "
          "grupo de 30 clientes apenas 2 das 20 instâncias foram resolvidas até a "
          "otimalidade comprovada, de modo que o valor do otimizador é, nas demais, apenas um "
          "limite superior. A Figura 3 ilustra a divergência entre os regimes de "
          "crescimento do tempo de processamento.")
    cur.figura("fig3_tempo_cpu.png")
    cur.legenda("Figura 3 – Tempo médio de processamento por número de clientes "
                "(escala logarítmica)", tam=10, depois=10)
    cur.p("Enquanto o tempo do método exato cresce de forma acentuada, de 2,7 a 6.542,1 "
          "segundos, aproximando-se do limite imposto, o tempo da meta-heurística "
          "evolui de maneira aproximadamente linear, de 11,8 a 67,5 segundos. Nenhuma "
          "execução ultrapassou 82 segundos em todo o conjunto S2, o que situa o serviço "
          "na faixa de tempo compatível com uso interativo.")

    cur.titulo("Efeito da distribuição espacial dos clientes", nivel=2)
    cur.p("O conjunto S2 divide-se em instâncias com clientes dispostos aleatoriamente e "
          "instâncias com clientes agrupados em aglomerados. A Tabela 4 compara o "
          "desempenho da meta-heurística nos dois subconjuntos, para os três cenários "
          "avaliados.")
    tab_random_clustered(cur)
    cur.p("No cenário completo, a meta-heurística apresenta desempenho relativo superior "
          "nas instâncias agrupadas (+0,27%) em comparação às aleatórias (+0,45%), ao "
          "mesmo tempo em que a aceleração praticamente dobra, passando de 52,9 para "
          "101,6 vezes. "
          "O resultado tem explicação direta: a estrutura em aglomerados agrava a "
          "dificuldade do método exato, que encerrou por limite de tempo 19 das 50 "
          "instâncias agrupadas contra 12 das 50 aleatórias, mas não penaliza de forma "
          "equivalente a busca local do esquema Ruin-and-Recreate, cuja estratégia de "
          "destruição radial é particularmente adequada a clientes espacialmente "
          "concentrados.")

    cur.titulo("Cenários sem simultaneidade e sem compartilhamento", nivel=2)
    cur.p("A Tabela 5 e a Figura 4 comparam o comportamento da meta-heurística nos três "
          "cenários avaliados sobre o conjunto S2.")
    tab_cenarios(cur)
    cur.figura("fig4_desvio_custo.png")
    cur.legenda("Figura 4 – Diferença média de custo por cenário e número de clientes",
                tam=10, depois=10)
    cur.p("O cenário sem compartilhamento é aquele em que a meta-heurística mais se "
          "aproxima da referência, com diferença média de apenas +0,16% e reprodução "
          "exata do ótimo em 76 das 100 instâncias. O resultado é coerente com a natureza "
          "do cenário: sem colaboração, o problema se decompõe em dois subproblemas "
          "independentes e substancialmente menores, cujo espaço de busca é mais "
          "facilmente varrido. Esse cenário funciona, portanto, como verificação de "
          "sanidade da implementação, atestando que os desvios observados nos demais "
          "casos decorrem da dificuldade do problema colaborativo e não de erro de "
          "modelagem.")
    cur.p("O cenário sem simultaneidade é o mais exigente dos três, com diferença média "
          "de +0,89% tanto no conjunto S1 quanto no S2. A degradação se concentra nas "
          "instâncias maiores: mantém-se em +0,17% no grupo de 10 clientes e sobe até "
          "+2,24% no grupo de 30. Ainda assim, a meta-heurística reproduziu exatamente o "
          "custo de referência em 57 das 100 instâncias do conjunto ampliado e o superou "
          "em três, todas com o método exato encerrado por limite de tempo.")
    cur.p("O comportamento é coerente com a natureza do cenário. Ao dispensar a visita "
          "única, o modelo admite que as duas transportadoras atendam separadamente o "
          "mesmo cliente e que uma delas o visite em duas passagens, o que amplia o "
          "espaço de soluções e exige da meta-heurística mecanismos que não são nativos "
          "do framework, notadamente o vínculo por transportadora entre as tarefas de "
          "uma mesma demanda. A busca faz uso efetivo dessa liberdade: em 26 das 112 "
          "instâncias a solução final atribui ao menos um cliente compartilhado às duas "
          "transportadoras, somando 35 clientes nessa condição, padrão que o cenário com "
          "simultaneidade proíbe por construção.")
    cur.p("O custo dessa generalidade aparece no tempo. O cenário sem simultaneidade "
          "exige, para cada cliente compartilhado, o dobro de tarefas do cenário "
          "restrito, e o tempo médio por instância sobe de 36,2 para 64,4 segundos no "
          "conjunto S2, com máximo de 212,6 segundos. A aceleração frente ao método "
          "exato cai proporcionalmente, mas permanece expressiva nos grupos maiores, "
          "chegando a 39,4 vezes no grupo de 30 clientes.")

    cur.titulo("Economia proporcionada pela colaboração", nivel=2)
    cur.p("O propósito último do modelo colaborativo é quantificar quanto se economiza "
          "quando duas transportadoras compartilham clientes, em comparação com a "
          "operação independente. A Tabela 6 e a Figura 5 apresentam essa economia, "
          "calculada tanto a partir das soluções do método exato quanto das soluções da "
          "meta-heurística.")
    tab_colaboracao(cur)
    cur.figura("fig5_economia_colaboracao.png")
    cur.legenda("Figura 5 – Economia da colaboração por grupo de instâncias: método exato "
                "e meta-heurística", tam=10, depois=10)
    cur.p("A colaboração horizontal proporcionou economia média de 9,50% no conjunto S2 e "
          "de 12,01% no conjunto S1, com máximo de 13,26% no grupo de 20 clientes. A "
          "economia é consistentemente maior nas instâncias aleatórias (11,74%) do que "
          "nas agrupadas (7,26%), o que se explica pela estrutura do problema: quando os "
          "aglomerados contêm clientes exclusivos de ambas as transportadoras, transferir "
          "os clientes compartilhados de uma para a outra não evita que ambas visitem o "
          "mesmo aglomerado, e o ganho da colaboração se reduz.")
    cur.p("O resultado de maior relevância prática, porém, está na última coluna da "
          "Tabela 6. A estimativa de economia produzida pela meta-heurística difere da "
          "obtida pelo método exato em, no máximo, 0,66 ponto percentual, e em apenas "
          "0,17 ponto percentual na média do conjunto S2. Isso significa que os erros "
          "cometidos pela meta-heurística nos dois cenários comparados são de magnitude "
          "semelhante e, portanto, cancelam-se em larga medida na razão que define a "
          "economia. A consequência é direta: para responder à pergunta gerencial "
          "“compensa colaborar, e quanto se ganha com isso?”, o serviço desenvolvido "
          "fornece, em dezenas de segundos, praticamente a mesma resposta que o método "
          "exato forneceria em horas.")
    cur.p("As Figuras 6 e 7 ilustram graficamente o efeito da colaboração sobre as rotas "
          "da instância 1, contrastando a operação independente com a operação "
          "colaborativa. Os depósitos das duas transportadoras estão assinalados por "
          "quadrados e os clientes por círculos; as cores indicam a transportadora "
          "responsável e os clientes compartilhados são representados pelas duas cores. "
          "Na operação independente, ambas as transportadoras percorrem os mesmos "
          "clientes compartilhados; na operação colaborativa, cada um deles é atendido "
          "por uma única transportadora, e as rotas se reorganizam de modo a encurtar o "
          "percurso agregado.")
    cur.figura("artigo_Im3.png", largura=Cm(11.5))
    cur.legenda("Figura 6 – Rotas da instância 1 sem colaboração (custo total 341,89)",
                tam=10, depois=10)
    cur.figura("artigo_Im4.png", largura=Cm(11.5))
    cur.legenda("Figura 7 – Rotas da instância 1 com colaboração (custo total 279,01)",
                tam=10, depois=4)
    cur.fonte_tabela("Nota: rotas obtidas pelo método exato. No cenário sem colaboração "
                     "a meta-heurística reproduziu exatamente o custo de 341,89; no cenário "
                     "colaborativo obteve 282,76, correspondente a desvio de 1,34%.")

    cur.titulo("Limitações e diferenças estruturais", nivel=2)
    cur.p("Apesar do esforço de espelhar fielmente o modelo matemático no framework "
          "Jsprit, subsistem diferenças estruturais que explicam os desvios observados e "
          "delimitam o alcance das conclusões. A mais evidente é de natureza: o otimizador "
          "exato retorna provas de otimalidade ou limites válidos, ao passo que a "
          "meta-heurística não oferece garantia alguma sobre a distância ao ótimo: as "
          "diferenças aqui reportadas são medidas empíricas, não cotas.")
    cur.p("A segunda diferença é de representação. O modelo exato emprega variáveis "
          "contínuas de fluxo que descrevem explicitamente a carga destinada a cada "
          "cliente em cada arco, enquanto o Jsprit rastreia apenas a carga agregada do "
          "veículo. A terceira, e mais consequente, diz respeito à alocação dos clientes "
          "compartilhados: no modelo exato essa decisão é endógena, e o otimizador explora "
          "simultaneamente todas as atribuições possíveis; no serviço desenvolvido ela é "
          "tratada pelo mecanismo de habilidades, que delega a decisão à busca sem "
          "enumerar exaustivamente as combinações. Para um número elevado de clientes "
          "compartilhados, algumas atribuições podem, portanto, nunca ser visitadas, o "
          "que ajuda a explicar a degradação observada nas instâncias de maior porte.")
    cur.p("Por fim, a comparação de tempos deve ser lida com a devida cautela. Embora "
          "ambos os métodos tenham sido executados na mesma configuração de hardware, o "
          "otimizador operou sob limite de 7.200 segundos, o que trunca seus tempos nas "
          "instâncias mais difíceis e, portanto, subestima a aceleração real da "
          "meta-heurística nesses casos. Em uma instância do grupo de 20 clientes, "
          "executada sem limite de tempo, o método exato consumiu 39.532,9 segundos, "
          "cerca de onze horas, contra 29,9 segundos da meta-heurística.")


def escrever_conclusao(cur):
    cur.p("Este trabalho desenvolveu e validou um serviço web para a resolução do "
          "problema de roteamento de veículos colaborativo com clientes compartilhados e "
          "coletas e entregas simultâneas. Foram implementados, sobre a biblioteca "
          "Jsprit, os quatro cenários do problema, o que exigiu traduzir para o paradigma "
          "orientado a objetos restrições originalmente formuladas de modo declarativo, "
          "notadamente o vínculo entre as tarefas de coleta e entrega de um mesmo cliente "
          "e o controle da alocação dos clientes compartilhados. A validação percorreu "
          "112 instâncias em cada um dos três cenários que dispõem de referência exata, "
          "totalizando 336 execuções comparadas com os resultados do método exato.")
    cur.p("Os resultados sustentam três conclusões principais. A primeira é que a "
          "meta-heurística aproxima o ótimo com fidelidade elevada no cenário completo do "
          "problema: diferença média de 0,22% no conjunto S1 e de 0,36% no conjunto S2, "
          "com reprodução exata do custo ótimo em 60 das 100 instâncias do conjunto "
          "ampliado e soluções superiores ao incumbente do otimizador em instâncias que este "
          "não resolveu até a otimalidade. A segunda é que o ganho computacional é de "
          "ordem de grandeza: nenhuma execução ultrapassou 82 segundos, contra médias "
          "superiores a 6.500 segundos do método exato nas maiores instâncias, com "
          "aceleração de até 96,9 vezes. Esses dois resultados, em conjunto, respondem "
          "afirmativamente à questão que motivou a pesquisa: é possível obter, em tempo "
          "compatível com uso interativo, soluções cuja distância ao ótimo não altera "
          "materialmente a decisão logística.")
    cur.p("A terceira conclusão é a de maior interesse aplicado. A economia proporcionada "
          "pela colaboração horizontal, estimada em 9,50% em média no conjunto S2, é "
          "reproduzida pela meta-heurística com desvio inferior a 0,2 ponto percentual em "
          "relação ao método exato. Como os erros da meta-heurística nos cenários "
          "colaborativo e independente têm magnitude semelhante e se cancelam na razão "
          "que define a economia, o serviço desenvolvido pode ser empregado com segurança "
          "para dimensionar o benefício de uma coalizão entre transportadoras, sem "
          "depender da disponibilidade de um otimizador comercial nem de horas de "
          "processamento.")
    cur.p("O trabalho também delimitou o alcance da abordagem. O cenário sem a "
          "restrição de simultaneidade, por admitir que um cliente compartilhado seja "
          "atendido separadamente pelas duas transportadoras, é o mais exigente: a "
          "diferença média sobe para 0,89% e atinge 2,24% nas instâncias de 30 clientes, "
          "e o tempo por instância quase dobra em relação ao cenário restrito. Em "
          "contrapartida, a meta-heurística faz uso efetivo dessa liberdade, atribuindo "
          "clientes compartilhados às duas transportadoras em 26 das 112 instâncias, o "
          "que confirma que a variante relaxada foi de fato explorada e não apenas "
          "aproximada pela solução do cenário restrito.")
    cur.p("Como desdobramentos, apontam-se quatro direções. A primeira é a implementação "
          "da variante muitos-para-muitos, na qual se elimina a figura do depósito "
          "central único e as demandas passam a ser expressas como pares origem-destino "
          "distribuídos arbitrariamente na rede. Trata-se de variante inicialmente prevista "
          "neste ciclo, cujo mapeamento das estruturas de dados foi iniciado, mas que não "
          "produziu resultados experimentais em tempo hábil e permanece, portanto, como "
          "trabalho futuro. A segunda é a calibração adaptativa dos parâmetros da "
          "meta-heurística em função do porte da instância, com atenção ao cenário sem "
          "simultaneidade, onde o espaço de busca é maior e a degradação se concentra "
          "nas instâncias de 30 clientes. A terceira é a adoção de estratégias de partida "
          "quente, inicializando a busca a partir de soluções construtivas de boa "
          "qualidade. A quarta é a extensão do modelo a janelas de tempo e frotas "
          "heterogêneas, ampliando sua aderência a contextos logísticos reais.")


REFERENCIAS = [
    "CACERES-CRUZ, J. et al. Rich vehicle routing problem: survey. ACM Computing "
    "Surveys, v. 47, n. 2, 2014.",
    "CORDEAU, J.-F.; LAPORTE, G.; ROPKE, S. Recent models and algorithms for one-to-one "
    "pickup and delivery problems. In: GOLDEN, B.; RAGHAVAN, S.; WASIL, E. (ed.). The "
    "Vehicle Routing Problem: latest advances and new challenges. Boston: Springer, "
    "2008. p. 327-357.",
    "FERNÁNDEZ, E.; ROCA-RIU, M.; SPERANZA, M. G. The shared customer collaboration "
    "vehicle routing problem. European Journal of Operational Research, v. 265, n. 3, "
    "p. 1078-1093, 2018.",
    "GANSTERER, M.; HARTL, R. F. Collaborative vehicle routing: a survey. European "
    "Journal of Operational Research, v. 268, n. 1, p. 1-12, 2018.",
    "GANSTERER, M.; HARTL, R. F. Shared resources in collaborative vehicle routing. TOP, "
    "v. 28, n. 1, p. 1-20, 2020.",
    "GRAPHHOPPER. jsprit: java based, open source toolkit for solving rich vehicle "
    "routing problems. Versão 1.9.0-beta.3. Disponível em: "
    "https://github.com/graphhopper/jsprit. Acesso em: 20 jul. 2026.",
    "GUROBI OPTIMIZATION. Gurobi Optimizer Reference Manual. Beaverton: Gurobi "
    "Optimization, 2024.",
    "HIMSTEDT, B.; MEISEL, F. A systematic evaluation of extensions for the shared "
    "customer collaboration vehicle routing problem. Bremen: Bundesvereinigung Logistik "
    "(BVL), 2021.",
    "PADMANABHAN, B. et al. Potential benefits of carrier collaboration in vehicle "
    "routing problem with pickup and delivery. Transportation Letters, v. 14, n. 3, "
    "p. 258-273, 2022.",
    "PAN, S. et al. Horizontal collaborative transport: survey of solutions and practical "
    "implementation issues. International Journal of Production Research, v. 57, "
    "n. 15-16, p. 5340-5361, 2019.",
    "PLOSKAS, N. et al. A tangible collaborative decision support system for various "
    "variants of the vehicle routing problem. Lecture Notes in Business Information "
    "Processing, v. 216, p. 73-84, 2015.",
    "RAN, L.; LI, L.; ZHAO, X. Brief review on heterogeneous vehicle routing problems. "
    "In: IEEE CHINESE CONTROL AND DECISION CONFERENCE. Anais [...]. IEEE, 2022. "
    "p. 5847-5852.",
    "SCHRIMPF, G. et al. Record breaking optimization results using the ruin and recreate "
    "principle. Journal of Computational Physics, v. 159, n. 2, p. 139-171, 2000.",
    "TORRES-RAMOS, A. F. et al. A GRASPxILS for the shared customer collaboration vehicle "
    "routing problem. IFAC-PapersOnLine, v. 52, n. 13, p. 2608-2613, 2019.",
    "VERDONCK, L. et al. Collaborative logistics from the perspective of road "
    "transportation companies. Transport Reviews, v. 33, n. 6, p. 700-719, 2013.",
    "WANG, Y. et al. Collaboration and transportation resource sharing in multiple "
    "centers vehicle routing optimization with delivery and pickup. Knowledge-Based "
    "Systems, v. 160, p. 296-310, 2018.",
    "ZHANG, W. et al. Composite multi-objective optimization on a new collaborative "
    "vehicle routing problem with shared carriers and depots. Journal of Cleaner "
    "Production, v. 274, 2020.",
]


# ======================================================================= main
def main():
    doc = Document(MODELO)
    remover_orientacoes(doc)
    preencher_capa(doc)
    p = doc.paragraphs

    # localiza as ancoras pelos titulos do modelo
    def ancora(titulo):
        for par in doc.paragraphs:
            if par.text.strip().upper().startswith(titulo):
                return par
        raise KeyError(titulo)

    secoes = [
        ("RESUMO", escrever_resumo),
        ("INTRODU", escrever_introducao),
        ("JUSTIFICATIVA", escrever_justificativa),
        ("OBJETIVOS", escrever_objetivos),
        ("METODOLOGIA", escrever_metodologia),
        ("RESULTADOS", escrever_resultados),
        ("CONCLUS", escrever_conclusao),
    ]
    for titulo, escritor in secoes:
        cab = ancora(titulo)
        # remove o paragrafo de instrucao do modelo, logo apos o titulo
        seguinte = cab._p.getnext()
        if seguinte is not None and seguinte.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph
            par_seg = Paragraph(seguinte, cab._parent)
            if par_seg.style.name in ("Normal", "Default") and par_seg.text.strip():
                remover(par_seg)
        escritor(Cursor(doc, cab))

    # referencias
    ref = ancora("REFER")
    seguinte = ref._p.getnext()
    from docx.text.paragraph import Paragraph
    while seguinte is not None and seguinte.tag.endswith("}p"):
        par_seg = Paragraph(seguinte, ref._parent)
        prox = seguinte.getnext()
        if par_seg.text.strip():
            remover(par_seg)
        seguinte = prox
    cur = Cursor(doc, ref)
    for entrada in REFERENCIAS:
        par = cur.p(entrada, recuo=False, alinh=WD_ALIGN_PARAGRAPH.LEFT)
        par.paragraph_format.line_spacing = 1
        par.paragraph_format.space_after = Pt(5)

    doc.save(SAIDA)
    print("Relatorio gerado:", SAIDA)


if __name__ == "__main__":
    main()
